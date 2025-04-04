from langchain_neo4j import Neo4jGraph
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LCDocument
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.documents import Document
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import streamlit as st
import numpy as np
import re
import os
import time

load_dotenv()


def read_docx(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])  # Extract all paragraphs
    return text


def read_pdf(file_path):
    """Reads and extracts text from a PDF file."""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"  # Extract text from each page

    # Clean the text to remove special tokens that might cause issues
    text = text.replace("<|endoftext|>", "")
    return text


def parse_best_practices(text):
    principles = []
    current_principle = None

    lines = text.strip().split("\n")  # Split by lines
    for line in lines:
        line = line.strip()

        principle_match = re.match(r"^Principle (\d+\.\d+): (.+)", line)
        practice_match = re.match(r"^Practice (\d+\.\d+): (.+)", line)

        if principle_match:
            if current_principle:
                principles.append(current_principle)
            current_principle = {
                "id": principle_match.group(1),
                "name": principle_match.group(2),
                "practices": [],
            }
        elif practice_match and current_principle:
            current_principle["practices"].append(
                {"id": practice_match.group(1), "description": practice_match.group(2)}
            )

    if current_principle:
        principles.append(current_principle)

    return principles


best_practices_file_path = "./data/The CyberGov™ Framework – Optimizing Your Cybersecurity Posture v. 8.0 14 Dec 2023.docx"
best_practices_text = read_docx(best_practices_file_path)
best_practices_doc = Document(
    best_practices_text, metadata={"source": best_practices_file_path}
)
principles = parse_best_practices(best_practices_text)


class ComplianceReport(BaseModel):
    status: str = Field(description="The compliance status: 'Pass' or 'Fail'")
    causality: str = Field(
        description="A to-the-point concise reason (Cause) for the compliance status (Effect)"
    )
    corrective_measures: str = Field(
        description="Suggested corrective measures if status is 'Fail', or empty if 'Pass'"
    )


# Initialize LangChain components
NEO4J_URL = os.getenv("NEO4J_URL")
NEO4J_USER = os.getenv("NEO4J_USER")
NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD")
graph = Neo4jGraph(url=NEO4J_URL, username=NEO4J_USER, password=NEO4J_PASSWORD)
embeddings = OpenAIEmbeddings(disallowed_special=())
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
llm = ChatOpenAI(
    model_name="gpt-4o-mini",
    temperature=0.1,
    logprobs=True,
    model_kwargs={"response_format": {"type": "json_object"}},
)
parser = PydanticOutputParser(pydantic_object=ComplianceReport)
prompt_template = PromptTemplate(
    template="""
You are an expert compliance analyst tasked with evaluating the compliance status of the best practice based on the provided context. 
The context consists of relevant remarks from board members. Clearly state the status in your response as "Pass" or "Fail" at the top.
You will be provided with a key indicator and a practice statement. You need to evaluate the compliance status of the practice based on the key indicator.

### Best Practice:
{practice}

### Key Indicator:
{key_indicator}

### Context:
{context}

### Question:
Based on the context, does the organization comply with this best practice? Provide reasoning if it doesn't and corrective measures. Your description should be easy to comprehend. If you don't find any relevant information, you can state that as well.

Format your response as a JSON object with the following fields:
- status: "Pass" or "Fail"
- causality: A to-the-point concise reason (Cause) for the compliance status (Effect)
- corrective_measures: Suggested actions if failing, or empty string if passing

### Answer:
""",
    input_variables=["practice", "key_indicator", "context"],
)
compliance_chain = prompt_template | llm


def check_compliance(practice_statement, key_indicator, vector_store, k=5):
    # Retrieve relevant documents
    retrieved_docs = vector_store.similarity_search_with_score(practice_statement, k=k)

    # Format retrieved documents into a structured context
    context = "\n\n".join(
        [
            f"Document {i+1} (score: {score}):\n{doc.page_content}"
            for i, (doc, score) in enumerate(retrieved_docs)
        ]
    )

    # Run the compliance chain
    result = compliance_chain.invoke(
        {
            "practice": practice_statement,
            "key_indicator": key_indicator,
            "context": context,
        }
    )
    return result

from PIL import Image
import streamlit as st
import base64

def get_base64_img(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()

image_base64 = get_base64_img("cybergov-icon.png")

# Streamlit App
def main():
    st.set_page_config(
        page_title="CyberGov",
        page_icon="./favicon.png",
        initial_sidebar_state="expanded",
        layout="wide"
    )
    # st.title("CyberGov Powered by Diogenes")
    st.markdown(
    f"""
    <div style='display: flex; align-items: center; justify-content: center; margin-top: 20px;'>
        <h3 style='display: flex; align-items: end; gap: 10px;'>
            <img src="data:image/png;base64,{image_base64}" style="height: 2em;" />
            Powered by Diogenes
        </h3>
    </div>
    """,
    unsafe_allow_html=True
)

    # Add sidebar for weights configuration
    with st.sidebar:
        st.title("Key Indicator Weights")
        st.markdown("Adjust the importance of each key indicator (in %)")
        
        # Create dictionary to store weights for each practice and its key indicators
        weights = {}
        weight_validation = {}
        
        # Initialize weights for all practices and their key indicators
        for principle in principles:
            for practice in principle['practices']:
                practice_id = practice["id"]
                result = graph.query(
                    """
                    MATCH (pr:Practice)-[:HAS_KEY_INDICATOR]->(ki:KeyIndicator)
                    WHERE pr.id = $practice_id
                    RETURN ki;
                    """,
                    params={"practice_id": practice_id},
                )
                
                if result:
                    # Store weights for each key indicator
                    if practice_id not in weights:
                        weights[practice_id] = {}
                        weight_validation[practice_id] = {"total": 0, "valid": False}
                    
                    # Count key indicators for this practice
                    ki_count = len(result)
                    default_weight = 100 // ki_count if ki_count > 0 else 100
                    remaining_weight = 100 - (default_weight * (ki_count - 1)) if ki_count > 0 else 0
                    
                    # Create an expander for each practice
                    with st.expander(f"Practice {practice_id}"):
                        total_weight = 0
                        for i, record in enumerate(result):
                            ki = record["ki"]
                            ki_id = ki["id"] if "id" in ki else f"{practice_id}-ki-{i}"
                            
                            # Last key indicator gets remaining weight to ensure total is 100
                            this_default = remaining_weight if i == ki_count - 1 else default_weight
                            
                            # Create slider for each key indicator
                            weight = st.slider(
                                f"KI: {ki['question'][:30]}...",
                                min_value=0,
                                max_value=100,
                                value=this_default,
                                key=f"weight_{practice_id}_{ki_id}"
                            )
                            weights[practice_id][ki_id] = weight
                            total_weight += weight
                        
                        # Show total weight
                        status = "OK" if total_weight == 100 else "should be 100%"
                        st.info(f"Total weight: {total_weight}% ({status})")
                        weight_validation[practice_id]["total"] = total_weight
                        weight_validation[practice_id]["valid"] = (total_weight == 100)
        
        # Add reset button to set equal weights
        if st.button("Reset to Equal Weights"):
            st.rerun()
        
        # Add submit button to check weights
        submit_disabled = not all(data["valid"] for data in weight_validation.values())
        submit_button = st.button("Submit Weights", disabled=submit_disabled)
        
        if submit_disabled:
            st.error("Please ensure all practice weights sum to 100% before submitting.")
        elif submit_button:
            st.success("Weights validated and saved successfully!")

    uploaded_file = st.file_uploader("Upload a DOCX or PDF file", type=["pdf", "docx"])

    if uploaded_file is not None:
        file_extension = uploaded_file.name.split(".")[-1].lower()

        if file_extension == "pdf":
            text = read_pdf(uploaded_file)
        elif file_extension == "docx":
            text = read_docx(uploaded_file)
        else:
            st.error("Unsupported file format.")
            return

        chunks = text_splitter.split_text(text)
        docs = [LCDocument(page_content=chunk) for chunk in chunks]
        vector_store_memo = FAISS.from_documents(docs, embeddings)

        for principle in principles:
            for practice in principle['practices']:
                ki_status = []
                ki_weights = []
                ki_conf_levels = []
                practice_id = practice["id"]
                result = graph.query(
                    """
                    MATCH (pr:Practice)-[:HAS_KEY_INDICATOR]->(ki:KeyIndicator)
                    WHERE pr.id = $practice_id
                    RETURN ki;
                    """,
                    params={"practice_id": practice_id},
                )

                # Create a practice-level container
                if result:
                    practice_container = st.container()
                    
                    # Display practice information in the container
                    with practice_container:
                        st.markdown(
                            f"#### Practice {practice['id']}: {practice['description']}"
                        )
                        
                        for i, record in enumerate(result):
                            key_indicator = record["ki"]

                            if key_indicator:
                                # Create a container for each key indicator set
                                ki_container = st.status(f"K{i+1}: {key_indicator['question']}", expanded=False, state='running')
                                
                                # Get key indicator ID for weight lookup
                                ki_id = key_indicator["id"] if "id" in key_indicator else f"{practice_id}-ki-{i}"
                                ki_weight = weights.get(practice_id, {}).get(ki_id, 100 // len(result))
                                
                                with ki_container:
                                    # st.markdown(f"**Key Indicator:** {key_indicator['question']} (Weight: {ki_weight}%)")

                                    # Create placeholders for the streaming effect
                                    status_placeholder = st.empty()
                                    confidence_placeholder = st.empty()
                                    cause_placeholder = st.empty()
                                    measures_placeholder = st.empty()

                                    # Get the compliance report
                                    report = check_compliance(
                                        practice["description"],
                                        key_indicator["question"],
                                        vector_store_memo,
                                    )

                                    # Extract confidence information
                                    confidence = ""
                                    for logprob in report.response_metadata["logprobs"]["content"]:
                                        if logprob["token"] == "Pass" or logprob["token"] == "Fail":
                                            confidence = np.exp(logprob['logprob'])*100

                                    # Parse the report
                                    report = parser.parse(report.content)
                                    ki_status.append(report.status)
                                    ki_weights.append(ki_weight)
                                    ki_conf_levels.append(confidence)

                                    # Display status with streaming effect
                                    
                                    # Status
                                    status_text = f"**Status:** {'🟢 ' if report.status == 'Pass' else '🔴 '}{report.status}"
                                    for i in range(len(status_text) + 1):
                                        status_placeholder.markdown(status_text[:i])
                                        time.sleep(0.001)
                                    
                                    # Confidence
                                    conf_text = f"**Confidence:** {confidence:.2f}%"
                                    for i in range(len(conf_text) + 1):
                                        confidence_placeholder.markdown(conf_text[:i])
                                        time.sleep(0.001)
                                    
                                    # Cause
                                    cause_text = f"**Cause:** {report.causality}"
                                    for i in range(len(cause_text) + 1):
                                        cause_placeholder.markdown(cause_text[:i])
                                        time.sleep(0.001)
                                    
                                    # Corrective measures
                                    measures_text = f"**Corrective measures:** {report.corrective_measures}"
                                    for i in range(len(measures_text) + 1):
                                        measures_placeholder.markdown(measures_text[:i])
                                        time.sleep(0.001)
                                    
                                    ki_container.update(state="complete") if report.status == "Pass" else ki_container.update(state="error")
                                    # Add a small separator between key indicators
                                    # st.write("---")
                        
                        # Display the overall compliance status for the practice
                        if ki_status:
                            # Calculate weighted average of confidence levels
                            fail_conf = sum(weight*conf_level for status, weight, conf_level in zip(ki_status, ki_weights, ki_conf_levels) if status == "Fail")
                            fail_percentage = fail_conf / 100
                            pass_percentage = 100 - fail_percentage
                            
                            
                            # Count traditional metrics as well
                            failed_count = sum(1 for status in ki_status if status == "Fail")
                            total_count = len(ki_status)
                            
                            # Determine if practice passes overall
                            # 80% threshold
                            threshold = pass_percentage >= 80
                            principle_status = "Pass" if threshold else "Fail"
                            
                            st.markdown(f"### Compliance Status")
                            st.markdown(f"**Status:** {'🟢 ' if principle_status == 'Pass' else '🔴 '}{principle_status}")
                            st.markdown(f"**Failed Indicators:** {failed_count}/{total_count}")
                            st.markdown(
                                """
                                <style>
                                    .stProgress > div > div > div > div {
                                        background-color: green;
                                    }
                                    .stProgress > div > div > div {
                                        background-color: red;
                                    }
                                </style>""",
                                unsafe_allow_html=True,
                            )
                            st.progress(pass_percentage/100, text=f"**Pass:** {pass_percentage:.1f}%")

                            # Separator between practices
                            st.markdown("---")

if __name__ == "__main__":
    main()
