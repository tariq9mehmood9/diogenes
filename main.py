import tkinter as tk
from tkinter import filedialog, scrolledtext
import docx
import os
from langchain_neo4j import Neo4jGraph
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LCDocument
from langchain.llms import OpenAI
from dotenv import load_dotenv
from pprint import pprint
from pydantic import BaseModel, Field
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.documents import Document
import json
import re
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

load_dotenv()


def read_docx(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])  # Extract all paragraphs
    return text


def read_pdf(file_path):
    """Reads and extracts text from a PDF file."""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"  # Extract text from each page
    return text


def parse_best_practices(text):
    principles = []
    current_principle = None

    lines = text.strip().split("\n")  # Split by lines
    for line in lines:
        line = line.strip()

        principle_match = re.match(r"^Principle (\d+\.\d+): (.+)", line)
        practice_match = re.match(r"^Practice (\d+\.\d+): (.+)", line)

        if principle_match:
            if current_principle:
                principles.append(current_principle)
            current_principle = {
                "id": principle_match.group(1),
                "name": principle_match.group(2),
                "practices": [],
            }
        elif practice_match and current_principle:
            current_principle["practices"].append(
                {"id": practice_match.group(1), "description": practice_match.group(2)}
            )

    if current_principle:
        principles.append(current_principle)

    return principles


best_practices_file_path = "./data/The CyberGov™ Framework – Optimizing Your Cybersecurity Posture v. 8.0 14 Dec 2023.docx"
best_practices_text = read_docx(best_practices_file_path)
best_practices_doc = Document(
    best_practices_text, metadata={"source": best_practices_file_path}
)
principles = parse_best_practices(best_practices_text)


class ComplianceReport(BaseModel):
    status: str = Field(description="The compliance status: 'Pass' or 'Fail'")
    causality: str = Field(
        description="A to-the-point concise reason (Cause) for the compliance status (Effect)"
    )
    corrective_measures: str = Field(
        description="Suggested corrective measures if status is 'Fail', or empty if 'Pass'"
    )


class DocxViewerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Content Viewer")
        self.root.geometry("800x600")

        # Create UI elements
        self.header_label = tk.Label(
            root, text="Document Content Viewer", font=("Arial", 16, "bold")
        )
        self.header_label.pack(pady=10)

        self.upload_button = tk.Button(
            root, text="Upload File (DOCX or PDF)", command=self.upload_file
        )
        self.upload_button.pack(pady=10)

        self.file_label = tk.Label(root, text="No file selected", font=("Arial", 10))
        self.file_label.pack(pady=5)

        self.content_area = scrolledtext.ScrolledText(
            root, wrap=tk.WORD, height=25, width=80
        )
        self.content_area.pack(pady=10, padx=20, fill=tk.BOTH, expand=True)

        # Initialize LangChain components
        self.graph = Neo4jGraph(
            url="bolt://localhost:7687", username="neo4j", password="password"
        )
        self.embeddings = OpenAIEmbeddings(disallowed_special=())
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=50
        )
        self.llm = ChatOpenAI(
            model_name="gpt-4o-mini",
            model_kwargs={"response_format": {"type": "json_object"}},
        )
        self.parser = PydanticOutputParser(pydantic_object=ComplianceReport)
        self.prompt_template = PromptTemplate(
            template="""
You are an expert compliance analyst tasked with evaluating the compliance status of the best practice based on the provided context. 
The context consists of relevant remarks from board members. Clearly state the status in your response as "Pass" or "Fail" at the top.
You will be provided with a key indicator and a practice statement. You need to evaluate the compliance status of the practice based on the key indicator.

### Best Practice:
{practice}

### Key Indicator:
{key_indicator}

### Context:
{context}

### Question:
Based on the context, does the organization comply with this best practice? Provide reasoning if it doesn't and corrective measures. Your description should be easy to comprehend. If you don't find any relevant information, you can state that as well.

Format your response as a JSON object with the following fields:
- status: "Pass" or "Fail"
- causality: A to-the-point concise reason (Cause) for the compliance status (Effect)
- corrective_measures: Suggested actions if failing, or empty string if passing

### Answer:
""",
            input_variables=["practice", "key_indicator", "context"],
        )
        self.compliance_chain = LLMChain(
            llm=self.llm, prompt=self.prompt_template, output_key="compliance_report"
        )

    def upload_file(self):
        file_path = filedialog.askopenfilename(
            title="Select a File",
            filetypes=[("Word Documents", "*.docx"), ("PDF Files", "*.pdf")],
        )

        if file_path:
            self.file_label.config(text=os.path.basename(file_path))
            if file_path.endswith(".docx"):
                self.parse_file(file_path, file_type="docx")
            elif file_path.endswith(".pdf"):
                self.parse_file(file_path, file_type="pdf")
            else:
                self.content_area.insert(tk.END, "Unsupported file type.\n")

    def parse_file(self, file_path, file_type):
        try:
            # Clear previous content
            self.content_area.delete(1.0, tk.END)

            # Read the file content
            if file_type == "docx":
                file_text = read_docx(file_path)
            elif file_type == "pdf":
                file_text = read_pdf(file_path)
            else:
                self.content_area.insert(tk.END, "Unsupported file type.\n")
                return

            # Clean the text to remove special tokens that might cause issues
            file_text = file_text.replace('<|endoftext|>', '')
            
            # Process the text with LangChain
            chunks = self.text_splitter.split_text(file_text)
            docs = [LCDocument(page_content=chunk) for chunk in chunks]
            vector_store_memo = FAISS.from_documents(docs, self.embeddings)

            for principle in principles:
                result = self.graph.query(
                    """
                    MATCH (p:Principle)-[:HAS_PRACTICE]->(pr:Practice)-[:HAS_KEY_INDICATOR]->(ki:KeyIndicator)
                    WHERE p.id = $principle_id
                    RETURN p, pr, ki;
                    """,
                    params={"principle_id": principle["id"]},
                )

                for record in result:
                    practice = record["pr"]
                    key_indicator = record["ki"]

                    if key_indicator:
                        self.content_area.insert(
                            tk.END, f"**Practice:** {practice['id']}\n"
                        )
                        self.content_area.insert(
                            tk.END, f"**Key Indicator:** {key_indicator['question']}\n"
                        )

                        report = self.check_compliance(
                            practice["description"],
                            key_indicator["question"],
                            vector_store_memo,
                        )
                        if report:
                            self.content_area.insert(
                                tk.END, f"**Status:** {report.status}\n"
                            )
                            self.content_area.insert(
                                tk.END, f"**Cause:** {report.causality}\n"
                            )
                            self.content_area.insert(
                                tk.END,
                                f"**Corrective measures:** {report.corrective_measures}\n\n",
                            )

        except Exception as e:
            self.content_area.insert(tk.END, f"Error parsing file: {str(e)}\n")

    def check_compliance(self, practice_statement, key_indicator, vector_store, k=5):
        # Retrieve relevant documents
        retrieved_docs = vector_store.similarity_search_with_score(
            practice_statement, k=k
        )

        # Format retrieved documents into a structured context
        context = "\n\n".join(
            [
                f"Document {i+1} (score: {score}):\n{doc.page_content}"
                for i, (doc, score) in enumerate(retrieved_docs)
            ]
        )

        # Run the compliance chain
        result = self.compliance_chain.invoke(
            {
                "practice": practice_statement,
                "key_indicator": key_indicator,
                "context": context,
            }
        )

        # Parse the JSON string into a ComplianceReport object
        try:
            json_str = result["compliance_report"]
            parsed_json = json.loads(json_str)
            return ComplianceReport(**parsed_json)
        except Exception as e:
            self.content_area.insert(
                tk.END,
                f"Failed to parse result into ComplianceReport model. Error: {e}\n",
            )
            self.content_area.insert(tk.END, "Raw result:\n")
            self.content_area.insert(tk.END, pprint(result["compliance_report"]))
            return None


if __name__ == "__main__":
    # Check if required libraries are installed
    try:
        import docx
        from PyPDF2 import PdfReader
    except ImportError as e:
        print(f"Required module not installed: {e}. Please install it using pip.")
        exit(1)

    root = tk.Tk()
    app = DocxViewerApp(root)
    root.mainloop()
